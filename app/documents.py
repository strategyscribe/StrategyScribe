"""Načítanie textových dokumentov (PDF, Word, txt/md) na doplnenie stratégie.

PDF sa neparsuje lokálne na text — posiela sa priamo na Claude API ako dokument
(Claude natívne číta PDF vrátane grafov a obrázkov, čo je pri trading
materiáloch podstatné). Veľké PDF sa delia na časti po stranách. Word a čisté
textové súbory sa čítajú lokálne a delia na textové časti.
"""

import io
from pathlib import Path

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")

# Claude API prijme max ~100 strán PDF na jedno volanie; menšie časti = nižšie
# riziko limitu veľkosti requestu a lepšia priebežná spätná väzba.
PDF_PAGES_PER_PART = 20
TEXT_CHARS_PER_PART = 15000


def load_document_parts(path):
    """Načíta dokument a vráti zoznam častí na analýzu. Každá časť je dvojica:
    ("pdf", bytes) — časť PDF súboru, posiela sa na Claude ako dokument, alebo
    ("text", str) — kus čistého textu.
    Pri nepodporovanom type alebo chybe čítania vyhodí RuntimeError."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf_parts(path)
    if suffix == ".docx":
        return _split_text(_read_docx_text(path))
    if suffix in (".txt", ".md"):
        return _split_text(_read_text_file(path))
    raise RuntimeError(
        f"Nepodporovaný typ súboru: {path.name} (podporované: {', '.join(SUPPORTED_EXTENSIONS)})"
    )


def _load_pdf_parts(path):
    from pypdf import PdfReader, PdfWriter

    try:
        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
    except Exception as exc:
        raise RuntimeError(f"Nepodarilo sa otvoriť PDF {path.name}: {exc}")
    if total_pages == 0:
        raise RuntimeError(f"PDF {path.name} neobsahuje žiadne strany.")

    parts = []
    for start in range(0, total_pages, PDF_PAGES_PER_PART):
        writer = PdfWriter()
        for page_index in range(start, min(start + PDF_PAGES_PER_PART, total_pages)):
            writer.add_page(reader.pages[page_index])
        buffer = io.BytesIO()
        writer.write(buffer)
        parts.append(("pdf", buffer.getvalue()))
    return parts


def _read_docx_text(path):
    from docx import Document

    try:
        document = Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"Nepodarilo sa otvoriť Word dokument {path.name}: {exc}")
    pieces = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                pieces.append(" | ".join(cells))
    text = "\n".join(pieces).strip()
    if not text:
        raise RuntimeError(f"Word dokument {path.name} neobsahuje žiadny text.")
    return text


def _read_text_file(path):
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1250", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError(f"Nepodarilo sa prečítať text zo súboru {path.name}.")
    text = text.strip()
    if not text:
        raise RuntimeError(f"Súbor {path.name} je prázdny.")
    return text


def _split_text(text):
    if len(text) <= TEXT_CHARS_PER_PART:
        return [("text", text)]
    parts = []
    remaining = text
    while remaining:
        if len(remaining) <= TEXT_CHARS_PER_PART:
            parts.append(("text", remaining))
            break
        # Delíme podľa možnosti na hranici odseku, nech sa netrhá veta.
        split_at = remaining.rfind("\n", TEXT_CHARS_PER_PART // 2, TEXT_CHARS_PER_PART)
        if split_at == -1:
            split_at = TEXT_CHARS_PER_PART
        parts.append(("text", remaining[:split_at].strip()))
        remaining = remaining[split_at:].strip()
    return [part for part in parts if part[1]]
