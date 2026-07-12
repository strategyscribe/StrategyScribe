"""Export textového zhrnutia (s jednoduchým markdown formátovaním od AI) do
skutočného Word (.docx) dokumentu."""

from docx import Document


def save_as_docx(text, output_path):
    doc = Document()
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(str(output_path))
